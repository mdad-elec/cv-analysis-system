import os
import re
from typing import Dict, List, Optional, Tuple
import pytesseract
from pdf2image import convert_from_path
from pypdf import PdfReader
import docx
import spacy
from PIL import Image
import cv2
import numpy as np
import traceback

from app.core.logging import logger
from app.models.documents import (
    CVDocument, DocumentStatus, DocumentType, ParsedCV, 
    PersonalInfo, Education, WorkExperience, 
    Skill, Project, Certification
)
from app.services.llm_service import LLMService

nlp = None

def get_nlp_model():
    global nlp
    if nlp is None:
        try:
            nlp = spacy.load("en_core_web_sm")
        except:
            spacy.cli.download("en_core_web_sm")
            nlp = spacy.load("en_core_web_sm")
    return nlp

class DocumentProcessor:
    def __init__(self):
        self.llm_service = LLMService()
    
    @staticmethod
    def _preprocess_text(text: str) -> str:
        text = text.replace('\u2022', '-')  # Replace bullet points
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
        text = re.sub(r'\s+', ' ', text).strip()  # Remove extra whitespace
        return text

    @staticmethod
    def _enhance_image_for_ocr(image):
        try:
            img_np = np.array(image)
            
            gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
            
            thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                          cv2.THRESH_BINARY, 21, 11)
            
            denoised = cv2.fastNlMeansDenoising(thresh, None, 10, 7, 21)
            
            kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
            sharpened = cv2.filter2D(denoised, -1, kernel)
            
            kernel = np.ones((1, 1), np.uint8)
            dilated = cv2.dilate(sharpened, kernel, iterations=1)
            
            return Image.fromarray(dilated)
        except Exception as e:
            logger.error(f"Image enhancement failed: {str(e)}\n{traceback.format_exc()}")
            return image
    
    @staticmethod
    def _extract_text_from_pdf(file_path: str) -> str:
        logger.info(f"Starting text extraction from PDF: {file_path}")
        
        direct_extraction_text = ""
        try:
            logger.info("Attempting direct text extraction with PdfReader")
            reader = PdfReader(file_path)
            logger.info(f"PDF has {len(reader.pages)} pages")
            
            page_texts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                page_texts.append(page_text)
                logger.info(f"Page {i+1}: Extracted {len(page_text)} characters")
            
            direct_extraction_text = "\n\n".join(page_texts)
            text_length = len(direct_extraction_text.strip())
            logger.info(f"Direct extraction yielded {text_length} characters")
            
            if text_length > 200: 
                logger.info("Direct extraction successful, returning text")
                return DocumentProcessor._preprocess_text(direct_extraction_text)
                
            logger.info(f"PDF text extraction returned limited content ({text_length} chars), falling back to OCR")
        except Exception as e:
            logger.warning(f"Error extracting text with PyPDF: {str(e)}\n{traceback.format_exc()}")
        
        try:
            logger.info("Starting OCR extraction process")
            logger.info(f"Converting PDF to images with DPI=300")
            images = convert_from_path(file_path, dpi=300, thread_count=os.cpu_count() or 4)
            logger.info(f"Converted PDF to {len(images)} images")
            
            extracted_text = []
            for i, img in enumerate(images):
                logger.info(f"Processing image {i+1}/{len(images)}")
                
                logger.info(f"Image size: {img.size}, mode: {img.mode}")
                
                logger.info("Enhancing image for OCR")
                enhanced_img = DocumentProcessor._enhance_image_for_ocr(img)
                
                logger.info("Applying OCR with pytesseract")
                page_text = pytesseract.image_to_string(
                    enhanced_img, 
                    config='--psm 6 --oem 3 -l eng+osd'
                )
                
                logger.info(f"OCR extracted {len(page_text)} characters from image {i+1}")
                extracted_text.append(page_text)
            
            ocr_text = "\n\n".join(extracted_text)
            ocr_text_length = len(ocr_text)
            logger.info(f"OCR extraction complete, yielded {ocr_text_length} characters")
            
            if ocr_text_length < len(direct_extraction_text.strip()) and len(direct_extraction_text.strip()) > 0:
                logger.info("Direct extraction yielded more text than OCR, using direct extraction")
                return DocumentProcessor._preprocess_text(direct_extraction_text)
            
            return DocumentProcessor._preprocess_text(ocr_text)
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}\n{traceback.format_exc()}")
            if direct_extraction_text:
                logger.info("Falling back to direct extraction result despite limited content")
                return DocumentProcessor._preprocess_text(direct_extraction_text)
            return ""
    
    @staticmethod
    def _extract_text_from_docx(file_path: str) -> str:
        logger.info(f"Starting text extraction from DOCX: {file_path}")
        try:
            doc = docx.Document(file_path)
            
            paragraphs = []
            
            logger.info(f"Document has {len(doc.paragraphs)} paragraphs")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append(para.text)
                    logger.info(f"Paragraph {i+1}: {len(para.text)} characters")
            
            logger.info(f"Document has {len(doc.tables)} tables")
            for i, table in enumerate(doc.tables):
                logger.info(f"Table {i+1} has {len(table.rows)} rows")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            text = "\n\n".join(paragraphs)
            logger.info(f"DOCX extraction complete, yielded {len(text)} characters")
            return DocumentProcessor._preprocess_text(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}\n{traceback.format_exc()}")
            return ""
            
    @staticmethod
    async def process_document(cv_document: CVDocument) -> Tuple[ParsedCV, Optional[str]]:
        """Process a document with improved LLM-based parsing."""
        logger.info(f"Processing document: {cv_document.file_path}, type: {cv_document.file_type}")
        try:
            processor = DocumentProcessor()
            
            raw_text = ""
            if cv_document.file_type == DocumentType.PDF:
                raw_text = DocumentProcessor._extract_text_from_pdf(cv_document.file_path)
            else:
                raw_text = DocumentProcessor._extract_text_from_docx(cv_document.file_path)
            
            logger.info(f"Extracted {len(raw_text)} characters from document")
            
            if not raw_text:
                logger.error("Failed to extract any text from document")
                return ParsedCV(raw_text=""), "Failed to extract text from document"
            
            # Initialize parsed CV with raw text
            parsed_cv = ParsedCV(raw_text=raw_text)
            
            # Extract basic information using NER for initial classification
            logger.info("Extracting basic information using NER")
            nlp_model = get_nlp_model()
            doc = nlp_model(raw_text[:5000])  # Process only first 5000 chars for efficiency
            
            # Extract personal info using traditional methods as fallback
            logger.info("Extracting personal info")
            parsed_cv.personal_info = DocumentProcessor._extract_personal_info(raw_text, doc)
            
            # Let LLM handle the complete parsing for all sections
            logger.info("Enhancing CV with LLM")
            enhanced_cv = await processor.llm_service.enhance_cv(parsed_cv)
            
            logger.info("Document processing complete")
            return enhanced_cv, None
        except Exception as e:
            error_message = f"Error processing document: {str(e)}"
            logger.error(f"{error_message}\n{traceback.format_exc()}")
            return ParsedCV(raw_text=raw_text if 'raw_text' in locals() else ""), error_message
        
    @staticmethod
    def _extract_personal_info(text: str, doc) -> PersonalInfo:
        """Extract basic personal information as an initial step."""
        personal_info = PersonalInfo()
        
        # Extract name using NER
        top_text = "\n".join(text.split("\n")[:10])  # Look at more lines
        person_entities = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        logger.info(f"Found {len(person_entities)} person entities: {person_entities}")
        
        # Extract email with robust pattern
        email_pattern = r'\b([a-zA-Z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            logger.info(f"Found email: {email_match.group(0)}")
        
        if person_entities:
            name = person_entities[0].strip()
            name = re.sub(r'[^a-zA-Z\s.-]', '', name).strip()  # Clean name
            if len(name.split()) >= 2 and len(name) > 3:  # Basic validation
                personal_info.name = name
                logger.info(f"Assigned name: {name}")
        
        if email_match:
            personal_info.email = email_match.group(0)
        
        phone_pattern = r'(?:\+\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            personal_info.phone = phone_match.group(0)
            logger.info(f"Found phone: {phone_match.group(0)}")
        
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/profile/|linkedin:)[\w\-./]+'
        github_pattern = r'(?:github\.com/|github:)[\w\-]+'
        linkedin_match = re.search(linkedin_pattern, text.lower())
        github_match = re.search(github_pattern, text.lower())
        
        if linkedin_match:
            personal_info.linkedin = linkedin_match.group(0)
            logger.info(f"Found LinkedIn: {linkedin_match.group(0)}")
        if github_match:
            personal_info.github = github_match.group(0)
            logger.info(f"Found GitHub: {github_match.group(0)}")
        
        return personal_info